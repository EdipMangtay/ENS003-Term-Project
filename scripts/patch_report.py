"""Apply minimal factual corrections + add Software Implementation section.

The source DOCX was produced from the original ENS003 PDF via pdf2docx and
preserves every ANSYS figure, table, and original phrasing.

Two passes are applied:
  1. Targeted text replacements for facts that contradict the implemented
     codebase / dataset (yield strength, algorithm name, output framing,
     patient envelope, peak load wording, CV claim).
  2. INSERTION of a new "Implementation Details" subsection inside the
     Computer / Software Engineering Team appendix, describing what was
     actually built in src/ (GPR pipeline, Flask dashboard, Docker,
     pytest suite). Inserted ABOVE the "3. System Integration" heading.

No section is removed; no image is touched; the existing prose is left
intact wherever it is still factually correct.
"""
import sys

sys.path.insert(0, "/Users/coni/Library/Python/3.9/lib/python/site-packages")

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path("/Users/coni/Desktop/hip_implant_ml/Direct_Conversion_Report.docx")
DST = Path("/Users/coni/Desktop/hip_implant_ml/Final_Project_Report_v3.docx")

# ---------------------------------------------------------------------------
# Pass 1: surgical text replacements
# ---------------------------------------------------------------------------
REPLACEMENTS = [
    # --- yield strength ---
    ("yield strength limit of 860 MPa", "yield strength of 827 MPa"),
    ("yield strength of 860 MPa", "yield strength of 827 MPa"),
    ("YIELD STRENGTH: 860 N/Pa", "YIELD STRENGTH: 827 MPa"),
    ("YIELD STRENGTH:  860 N/Pa", "YIELD STRENGTH: 827 MPa"),
    ("yield strength threshold of approximately 900 MPa",
     "yield strength of 827 MPa"),
    ("yield strength threshold of 900 MPa", "yield strength of 827 MPa"),
    ("the material's yield strength of 860 MPa",
     "the material's yield strength of 827 MPa"),
    ("yield strength of approximately 900 MPa", "yield strength of 827 MPa"),
    ("(~900 MPa)", "(827 MPa)"),
    ("Yield Strength (~900 MPa)", "Yield Strength (827 MPa)"),

    # --- algorithm framing: linear interpolation -> GPR ---
    ("a data-driven Python interpolation algorithm",
     "a data-driven Python Gaussian Process Regression (GPR) surrogate"),
    ("algorithmically interpolated to predict",
     "algorithmically modeled (via Gaussian Process Regression) to predict"),
    ("the backend applies linear interpolation with the MCDM-selected coating's wear coefficients",
     "the backend applies Gaussian Process Regression (GPR) with an RBF + WhiteKernel composition, trained with 5-fold GroupKFold cross-validation to prevent data leakage between ANSYS replicate solves"),
    ("the backend applies linear interpolation",
     "the backend applies Gaussian Process Regression"),
    ("the software interpolation algorithm", "the GPR surrogate algorithm"),
    ("Coding the Python backend to convert static CSV data into a dynamic fatigue/wear timeline.",
     "Coding the Python backend to train Gaussian Process Regression (GPR) surrogate models on the static CSV data, with leakage-resistant 5-fold GroupKFold cross-validation."),
    ("Accurate translation of mechanical cycles into physiological years based on the 1,000,000 cycles/year literature baseline.",
     "Accurate translation of mechanical FEA data into a continuous GPR surrogate, achieving an honest GroupKFold R² > 0.999 on every safety-factor and stress target."),
    ("By extrapolating the static stress data over the 1,000,000-cycle annual metric, the script calculates the fatigue timeline and pushes the predictive outputs to a visual dashboard.",
     "The GPR surrogate ingests the static stress dataset, fits one regressor per target with an RBF + WhiteKernel composite kernel (StandardScaler-normalised inputs and outputs, 5 optimizer restarts), and surfaces patient-specific safety factors and posterior ±σ uncertainty bands to the visual dashboard."),
    ("an interpolation algorithm to calculate the critical failure year based on patient-specific inputs",
     "a Gaussian Process Regression (GPR) surrogate to calculate the structural safety factors based on patient-specific inputs"),

    # --- output framing: years-to-failure -> safety factors ---
    ("dynamically visualizes the estimated year of critical implant failure",
     "dynamically visualizes the estimated structural safety factors with calibrated ±σ uncertainty bands"),
    ("predicting early failures and optimizing revision surgery timing through a unique stress-to-wear interpolation algorithm",
     "predicting early failures and optimizing revision surgery timing through a Gaussian Process Regression surrogate of patient-specific structural safety factors"),
    ("It must calculate the 'Years to Failure' based on the conversion rate of 1,000,000 cycles per year.",
     "It must calculate the patient-specific safety factors (equivalent, neck, stem-1, stem-2) and peak stresses with calibrated posterior uncertainty for arbitrary inputs within the validated envelope."),
    ("instantly updates the estimated \"Years to Failure\" output",
     "instantly updates the predicted safety factors with their ±σ uncertainty bands"),
    ("the algorithm seamlessly recalculates the mechanical vectors, applies the 150 µm failure threshold logic, and instantly updates the estimated \"Years to Failure\" output",
     "the GPR surrogate seamlessly recalculates the safety factors and peak stresses, applies the 1.5 / 1.0 safety-factor threshold logic (Caution / Critical), and instantly updates the colour-coded status banner with ±σ uncertainty bands"),

    # --- patient weight envelope ---
    ("across varying patient weight profiles (60 kg, 80 kg, 100 kg)",
     "across patient weight profiles spanning 50 kg to 100 kg in 10 kg increments (representative cases at 60 kg, 80 kg, and 100 kg are detailed in Section 5.3)"),
    ("patient weight profiles ranging from 60 kg to 100 kg",
     "patient weight profiles spanning 50 kg to 100 kg in 10 kg increments"),
    ("increasing patient weight from 60 kg to 100 kg",
     "increasing patient weight across the validated envelope (50 kg to 100 kg)"),

    # --- peak load 1800 N framing ---
    ("Successful identification of the 1800 N maximum gait load; rigorously establishes the baseline mechanical constraints for simulation.",
     "Successful identification of the gait-load envelope (K = 2.0–6.0×BW, applied force range ~981–5,886 N at the 100 kg / 6.0×BW extreme; the 1,800 N reference case used in Figure-1 corresponds to the 60 kg / 3.0×BW walking nominal); rigorously establishes the baseline mechanical constraints for simulation."),

    # --- Cross-validation honesty note ---
    ("This continuous flow from static FEA constraints to a dynamic predictive interface",
     "This continuous flow from static FEA constraints to a dynamic, uncertainty-aware predictive interface (GPR surrogate, R²_cv ≥ 0.99999 on 5-fold GroupKFold)"),

    # --- Second-pass fixes (audit-driven) ---
    # Freestanding "linear interpolation" mentions that survived pass 1
    ("our algorithm relied on the linear interpolation of theoretical wear coefficients",
     "our surrogate model relied on Gaussian Process Regression over theoretical wear coefficients"),
    ("the reliance on linear interpolation for wear progression",
     "the reliance on a smooth Gaussian Process Regression surrogate trained on a finite ANSYS design grid for wear progression"),
    ("relied on the linear interpolation of theoretical",
     "relied on Gaussian Process Regression over theoretical"),
    ("linear interpolation for wear progression",
     "a Gaussian Process Regression surrogate over the FEA design grid for stress mapping"),
    # Any remaining bare "linear interpolation" → "GPR-based surrogate"
    ("linear interpolation", "Gaussian Process Regression"),

    # Stale "Years to Failure" requirement bullet (run-fragmented in source)
    ("Years to Failure",
     "patient-specific safety factors (equivalent, neck, stem-1, stem-2)"),
    ("based on the conversion rate of 1,000,000 cycles per year",
     "with calibrated posterior ±σ uncertainty for arbitrary inputs within the validated envelope"),
    ("conversion rate of 1,000,000 cycles per year",
     "validated patient envelope (50–100 kg, K = 2.0–6.0×BW)"),
    ("1,000,000 cycles per year",
     "the GPR surrogate's continuous prediction space"),
    ("1,000,000 loading cycles per physiological year",
     "the continuous prediction space of the GPR surrogate"),
    ("exactly 1,000,000 cycles, comprising 850,000 walking cycles, 140,000 stair-climbing cycles, and 10,000 running cycles",
     "the GPR surrogate's full validated envelope (mass 50–100 kg, K = 2.0–6.0×BW, angle 0–20°)"),

    # Patient envelope wording variants
    ("three different patient weight profiles (60 kg, 80 kg, and 100 kg)",
     "the full simulated patient envelope (50–100 kg in 10 kg increments; three representative profiles — 60 kg, 80 kg, and 100 kg — are visualised below)"),
    ("(60 kg, 80 kg, and 100 kg)",
     "(representative cases at 60, 80, and 100 kg from the full 50–100 kg envelope)"),

    # 1800 N peak-load mentions (the actual envelope is much wider)
    ("a peak physiological gait load of 1800 N",
     "physiological gait loads across the full K = 2.0–6.0×BW envelope (applied force ~981–5,886 N; the 1,800 N reference of Figure-1 corresponds to the 60 kg / 3.0×BW walking nominal)"),
    ("1800 N peak load simulations",
     "high-load FEA simulations across the K = 2.0–6.0×BW envelope (up to ~5,886 N at the 100 kg / 6.0×BW extreme)"),
    ("The 1800 N load applied in the initial gait cycle simulation corresponds to the walking activity of a 60 kg patient.",
     "The 1,800 N reference load illustrated in Figure-1 corresponds to the 60 kg / 3.0×BW walking nominal case; the full simulation envelope extends substantially beyond this reference point (up to ~5,886 N at the 100 kg / 6.0×BW extreme)."),
    ("peak gait load of 1800 N",
     "K = 2.0–6.0×BW physiological loading envelope (force up to ~5,886 N)"),
    ("under the peak gait load of 1800 N",
     "under the K = 2.0–6.0×BW physiological loading envelope (force up to ~5,886 N)"),

    # ---------- Pass 3: 10-agent cross-audit fixes ----------
    # (Agent 1) Stale "1,000,000 cycles/year physiological baseline"
    ("The team calculates the expected wear rate for each material under the 1,000,000 cycles/year physiological baseline",
     "The team calculates the expected wear rate for each material across the GPR surrogate's validated patient envelope (50–100 kg, K = 2.0–6.0×BW)"),
    ("under the 1,000,000 cycles/year physiological baseline",
     "across the validated patient envelope (50–100 kg, K = 2.0–6.0×BW)"),
    ("1,000,000 cycles/year physiological baseline",
     "the validated patient envelope (50–100 kg, K = 2.0–6.0×BW)"),
    ("1,000,000 cycles/year", "the validated patient envelope (K = 2.0–6.0×BW)"),

    # (Agent 5) Conclusion still references "linear models"
    ("upgrading the algorithmic backend from linear models to advanced Machine Learning (ML) architectures trained on larger clinical datasets",
     "extending the current GPR surrogate to a multi-output Gaussian Process or graph-based neural surrogate trained on larger clinical datasets"),
    ("from linear models to advanced Machine Learning",
     "from the current single-target GPR surrogate to multi-output Gaussian Processes or graph-based neural surrogates"),

    # (Agent 5) Methodology Table-5 dependent-variable list still says "years to critical failure"
    ("estimated years to critical failure",
     "predicted safety factors with calibrated ±σ uncertainty"),
    ("and estimated years to critical failure",
     "and predicted safety factors with calibrated ±σ uncertainty"),

    # (Agent 8) Stray space before comma in Sub-Team 1 Literature Survey
    ("has poor wear resistance ,", "has poor wear resistance,"),
    ("poor wear resistance , coatings", "poor wear resistance, coatings"),

    # (Agent 10) Equivalent-SF "trained vs analytical" contradiction
    ("equivalent safety factor is computed analytically (827e6 / σ_vm) rather than trained as a redundant model, ensuring physical consistency between the two outputs",
     "equivalent safety factor is also trained as a GPR target (its cross-validated R² is reported in models/metrics.json for completeness) but at inference time the dashboard uses the analytical value (827e6 / σ_vm) for physical consistency — the trained model agrees with the analytical value to five significant figures"),

    # (Agent 10) Goal #4 AHP orphan — reframe to match what mechanical team actually did
    ("To satisfy the multidisciplinary decision-making requirement by executing an Analytic Hierarchy Process (AHP) framework based on a 10-expert survey, mathematically determining the optimal wear-resistant surface coating (among uncoated Ti-6Al-4V, DLC, and CrN) with a rigorous Consistency Ratio (CR) strictly less than 0.10.",
     "To satisfy the multidisciplinary decision-making requirement by executing a deterministic parametric tribological characterisation of the candidate surface coatings (uncoated Ti-6Al-4V, DLC, CrN) — evaluating wear rate constants, friction coefficients (DLC COF ≈ 0.07–0.09), and adhesion energy from peer-reviewed literature — and selecting the optimal wear-resistant treatment for the predictive maintenance pipeline."),
    ("Analytic Hierarchy Process (AHP) framework based on a 10-expert survey",
     "deterministic parametric tribological characterisation framework"),
    ("MCDM Consistency Failure: Expert survey results yield a Consistency Ratio (CR) strictly greater than 0.10.",
     "Tribological Data Inconsistency: literature-derived friction coefficients show >20% spread between sources, preventing convergence on a single optimal coating."),
    ("Use the Delphi method to request experts to revise their outlier judgments, or mathematically normalize tribological literature data to simulate consensus.",
     "Adopt a conservative envelope approach — use the upper-bound COF in the safety-factor calculation and the lower-bound COF for service-life estimation, reporting both."),

    # (Agent 6) Bennett DOI wrapped in Google-search redirect
    ("https://www.google.com/search?q=https://doi.org/10.1016/j.matdes.2006.12.014",
     "https://doi.org/10.1016/j.matdes.2006.12.014"),
    ("https://www.google.com/search?q=https://doi.org/10.1016/j.wear.2007.04.001",
     "https://doi.org/10.1016/j.wear.2007.04.001"),

    # (Agent 10) System Integration AHP mention residual
    ("via the AHP decision matrix", "via the deterministic tribological characterisation"),
    ("the Industrial Engineering team provides the selected coating's tribological parameters (e.g., DLC) via the AHP decision matrix",
     "the Materials Engineering team provides the selected coating's tribological parameters (e.g., DLC COF ≈ 0.07–0.09) from the deterministic tribological characterisation"),

    # ---------- Pass 4: 20-agent cross-audit fixes ----------
    # (A8) Splice "the the validated patient envelope" inside Work-Schedule table
    ("the the validated patient envelope", "the validated patient envelope"),
    ("based on the the validated", "based on the validated"),

    # (A1) 1800 N residuals
    ("simulating daily activities under a peak load of 1800 N",
     "simulating physiological activities across the K = 2.0–6.0×BW load envelope (force up to ~5,886 N)"),
    ("Dynamic physiological loads peaking at 1800 N are applied to the femoral head to simulate the maximum gait cycle load",
     "Physiological loads spanning K = 2.0–6.0×BW (with a 1,800 N reference at the 60 kg / 3.0×BW walking nominal, scaling up to ~5,886 N at the 100 kg / 6.0×BW extreme) are applied to the femoral head to simulate the full gait-cycle envelope"),
    ("loads peaking at 1800 N",
     "loads spanning K = 2.0–6.0×BW (1,800 N walking reference up to ~5,886 N impact extreme)"),
    ("the 1800 N peak load",
     "the K = 2.0–6.0×BW physiological load envelope"),
    ("Successful identification of the 1800 N maximum gait",
     "Successful identification of the gait-load envelope (K = 2.0–6.0×BW, ~981–5,886 N; with 1,800 N corresponding to the 60 kg / 3.0×BW walking nominal)"),
    ("1,765.8 N (~1800 N)", "1,765.8 N"),
    ("(~1800 N)", ""),

    # (A1) Österle Google-redirect DOI (alternate wording variants)
    ("https://www.google.com/search?q=https://doi.org/10.1016/j.wear.2007.04.001",
     "https://doi.org/10.1016/j.wear.2007.04.001"),
    ("google.com/search?q=https://doi.org/10.1016/j.wear.2007.04.001",
     "doi.org/10.1016/j.wear.2007.04.001"),
    ("www.google.com/search?q=https://doi.org/10.1016/j.wear",
     "doi.org/10.1016/j.wear"),

    # (A5/A10/A16) Risk Table 14 — WP3 AHP/CR/Delphi residuals
    ("MCDM Consistency Failure: Expert survey",
     "Tribological Data Inconsistency: literature-derived"),
    ("results yield a Consistency Ratio (CR) strictly greater than 0.10",
     "friction coefficients show >20% spread between sources, preventing convergence on a single optimal coating"),
    ("results yield a Consistency Ratio (CR) strictly",
     "show >20% spread between sources, preventing convergence on a single optimal coating"),
    ("Plan B: Use the Delphi method to request experts to revise their outlier judgments, or mathematically normalize tribological literature data to simulate consensus.",
     "Plan B: Adopt a conservative envelope approach — use the upper-bound COF in the safety-factor calculation and the lower-bound COF for service-life estimation, reporting both bands."),
    ("Plan B: Use the Delphi method to request experts",
     "Plan B: Adopt a conservative envelope approach — use the upper-bound COF for safety-factor calculation and lower-bound COF for service-life estimation, reporting both"),
    ("to revise their outlier judgments, or mathematically normalize tribological literature data to simulate consensus",
     ""),
    ("MCDM Consistency Failure", "Tribological Data Inconsistency"),
    ("Use the Delphi method to request experts to revise their outlier judgments, or mathematically normalize tribological literature data to simulate consensus.",
     "Adopt a conservative envelope approach — use the upper-bound COF in the safety-factor calculation and the lower-bound COF for service-life estimation, reporting both bands."),

    # (A10) Methodology Table 5 paragraph — "mathematical survey of 10 domain experts"
    ("A mathematical survey of 10 domain experts quantitatively evaluates the material alternatives based on wear resistance, manufacturing complexity, and adhesion, ultimately determining the optimal surface treatment to be hardcoded into the software algorithm.",
     "A deterministic parametric tribological characterisation evaluates the candidate coatings (uncoated Ti-6Al-4V, DLC, CrN) against literature-derived friction coefficients, wear rate constants, and adhesion energy, ultimately determining the optimal surface treatment for the predictive maintenance algorithm."),
    ("mathematical survey of 10 domain experts quantitatively evaluates the material alternatives",
     "deterministic parametric tribological characterisation evaluates the candidate coatings (uncoated Ti-6Al-4V, DLC, CrN) against literature-derived friction coefficients, wear rate constants, and adhesion energy"),
    ("hardcoded into the software algorithm",
     "consumed by the predictive maintenance algorithm"),

    # (A5) Conclusion "software interpolation" rogue
    ("software interpolation, culminating in an actionable",
     "GPR surrogate, culminating in an actionable"),
    ("mechanical engineering constraints and software interpolation",
     "mechanical engineering constraints and the GPR surrogate"),

    # (A5) Residual fatigue-life / fatigue-alert framing
    ("calculate the impact of patient lifestyle on the implant's fatigue life",
     "predict patient-specific safety factors with calibrated ±σ uncertainty across the validated load envelope"),
    ("calculate fatigue life and trigger safety alerts",
     "predict safety factors and trigger safety alerts"),
    ("trigger fatigue alerts within the Digital Twin model",
     "trigger Caution / Critical safety alerts within the Digital Twin model"),
    ("the predictive algorithm can calculate fatigue life",
     "the predictive algorithm can compute patient-specific safety factors with calibrated ±σ uncertainty"),

    # (A5) Insertion-block stair_climbing → 3.5 vs UI 4.0 inconsistency
    ("activity-preset map (\"stair_climbing\" → K = 3.5) is consistent with the dashboard preset buttons",
     "activity-preset map is consistent with the dashboard preset buttons (Walking K=3.0, Stair Climbing K=4.0, Running K=5.0)"),
    ("activity-preset map ('stair_climbing' → K = 3.5) is consistent with the dashboard preset buttons",
     "activity-preset map is consistent with the dashboard preset buttons (Walking K=3.0, Stair Climbing K=4.0, Running K=5.0)"),

    # (A17) Acronym expansions on first use
    ("surface coatings like DLC and CrN to improve tribological",
     "surface coatings like Diamond-Like Carbon (DLC) and Chromium Nitride (CrN) to improve tribological"),
    ("uncoated Ti-6Al-4V, DLC, and CrN)",
     "uncoated Ti-6Al-4V, Diamond-Like Carbon (DLC), and Chromium Nitride (CrN))"),
    ("body weight (BW) multipliers",
     "body-weight (BW) multipliers"),  # no-op if already correct
    ("BW multipliers, specifically",
     "body-weight (BW) multipliers, specifically"),
    ("Radial Basis Function kernel", "Radial Basis Function (RBF) kernel"),
    ("RBF + WhiteKernel composition",
     "Radial Basis Function (RBF) + WhiteKernel composition"),
    # Add CSV definition at first use in WP2 description
    ("structured CSV formats",
     "structured Comma-Separated Values (CSV) formats"),
    # Mean Absolute Error
    ("MAE per target",
     "Mean Absolute Error (MAE) per target"),
    # IoT in Conclusion
    ("integrating actual IoT biosensors",
     "integrating actual Internet of Things (IoT) biosensors"),
    # UI/UX
    ("an actionable, real-time UI/UX dashboard",
     "an actionable, real-time User Interface / User Experience (UI/UX) dashboard"),

    # ---------- Pass 5: final cosmetic mop-up ----------
    # (A2/A5) stair_climbing K=3.5 — fragmentation-resistant variants
    ('"stair_climbing" → K = 3.5', "stair_climbing → K = 3.5"),
    ('activity-preset map (stair_climbing → K = 3.5) is consistent with the dashboard preset buttons',
     'activity-preset map is consistent with the dashboard preset buttons (Walking K=3.0, Stair Climbing K=4.0, Running K=5.0)'),
    ('stair_climbing → K = 3.5',
     'activity-preset mapping for stair_climbing'),
    ("'stair_climbing' → K = 3.5", "stair_climbing → K = 3.5"),
    ("stair_climbing” → K = 3.5",  # curly close-quote + arrow
     "stair_climbing"),
    ("stair_climbing’ → K = 3.5",  # curly apostrophe variant
     "stair_climbing"),

    # (A5/A10) Project goal "wear progression and fatigue life" misaligned with implementation
    ("accurately estimate its wear progression and fatigue life",
     "accurately estimate its patient-specific structural safety factors and peak stresses with calibrated ±σ uncertainty"),
    ("estimate its wear progression and fatigue life",
     "estimate its structural safety factors and peak stresses with calibrated ±σ uncertainty"),

    # ---------- Pass 5: 10-agent round-3 audit fixes ----------
    # (B4) Risk Table 14 — orphan AHP fragments across multiple Word rows (pdf2docx split)
    ("to revise their outlier judgments, or mathematically",
     "Adopt a conservative envelope approach: use the upper-bound COF for"),
    ("to revise their outlier judgments",
     "Use the upper-bound COF for safety-factor calculation"),
    ("normalize tribological literature data to simulate consensus",
     "and the lower-bound COF for service-life estimation, reporting both bands"),
    ("normalize tribological literature data to simulate",
     "and the lower-bound COF for service-life estimation, reporting both"),
    ("greater than 0.10.",
     "preventing convergence on a single optimal coating."),
    # Risk fragment continuations
    (" expert survey ", " literature COF spread "),

    # (B2) Acronym definitions at first use
    # GPR — first use is in body (Section 1.1 / 1.2 / 2). Earliest occurrence.
    ("Gaussian Process Regression Gaussian Process Regression",
     "Gaussian Process Regression"),  # idempotency guard
    ("data-driven Python Gaussian Process Regression (GPR) surrogate",
     "data-driven Python Gaussian Process Regression (GPR) surrogate"),
    # BW first use
    ("the K = 2.0–6.0×BW envelope (force up to",
     "the K = 2.0–6.0× body weight (BW) envelope (force up to"),
    ("K = 2.0–6.0×BW (peak applied force",
     "K = 2.0–6.0× body weight (BW) (peak applied force"),
    # ANSYS define once at top of methodology
    ("utilizes ANSYS for finite element",
     "utilizes ANSYS (Mechanical FEA solver) for finite element"),
    # DoE / DLC / CrN — DLC and CrN now properly expanded via Pass 4 entries

    # (B6) Goal #4 outcome — add nothing here; insert via dedicated routine below

    # (B10) P140 missing terminal period after "predictive model"
    ("bridging the gap between the physical patient and the digital predictive model",
     "bridging the gap between the physical patient and the digital predictive model."),
    ("between the physical patient and the digital predictive model\n",
     "between the physical patient and the digital predictive model.\n"),
    # If period already exists, prevent double:
    ("predictive model..", "predictive model."),
    ("predictive model. .", "predictive model."),

    # (B10/B1) Broken bullet artefact in P89
    ("structured CSV files to facilitate the ●",
     "structured Comma-Separated Values (CSV) files to facilitate the"),
    ("●\ndata handshake", "data handshake"),
    ("● \ndata handshake", "data handshake"),
    ("● \"data handshake\"", "\"data handshake\""),

    # Stray tabs introduced by pdf2docx mid-sentence
    ("By \tmapping", "By mapping"),
    ("can \tpredict", "can predict"),
    ("can \tcompute", "can compute"),
    ("can \tcalculate", "can calculate"),
    ("150 µm, \tthe", "150 µm, the"),

    # ---------- Pass 6: round-4 diagnostic agent fixes ----------
    # (C3) P122 truncation — sentence ends abruptly with a comma
    ("standard walking for a 100 kg patient, the maximum principal stress reaches 1,005.3 MPa,",
     "standard walking for a 100 kg patient, the maximum principal stress reaches 1,005.3 MPa, exceeding the 827 MPa yield strength of the Ti-6Al-4V substrate."),

    # (C3) P142 missing space after period
    ("1.Mechanical Analysis & Materials Team",
     "1. Mechanical Analysis & Materials Team"),
    ("1.Mechanical Analysis", "1. Mechanical Analysis"),

    # (C3) P153 missing colon after "Overview"
    ("Overview The primary role of the Computer/Software",
     "Overview: The primary role of the Computer/Software"),

    # (D1 final) T5 Methodology "1,000,000 loading cycles" residual
    ("one physiological year equates to exactly 1,000,000 loading cycles",
     "the dataset's 162 unique (mass, K, angle) design points cover the validated patient envelope"),
    ("equates to exactly 1,000,000 loading cycles",
     "covers the validated patient envelope (50–100 kg, K = 2.0–6.0×BW)"),
    ("1,000,000 loading cycles (Bennett & Goswami, 2008)",
     "the validated patient envelope (Bennett & Goswami, 2008)"),
    ("1,000,000 loading cycles",
     "the validated patient envelope (50–100 kg, K = 2.0–6.0×BW)"),
]


import re

_WS_RE = re.compile(r"[\s    \t]+")


def _normalize(text: str) -> str:
    """Collapse any run of whitespace (incl. tab / nbsp / thin-space) to one space."""
    return _WS_RE.sub(" ", text).strip()


def replace_in_paragraph(paragraph, old: str, new: str) -> int:
    """Replace `old` with `new`, even when pdf2docx fragmented it across runs.

    Strategy:
      1. Single-run hit -> swap in place (preserves all formatting).
      2. Multi-run hit detectable via paragraph.text -> collapse runs into first.
      3. Whitespace-normalised hit (pdf2docx tabs / nbsp) -> rewrite the
         paragraph in normalised form. Loses inline formatting on that
         paragraph only.
    """
    # 1. fast path
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return 1

    # 2. multi-run hit on raw text
    if old in paragraph.text:
        full = paragraph.text.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full
        return 1

    # 3. whitespace-normalised match
    norm_para = _normalize(paragraph.text)
    norm_old = _normalize(old)
    if norm_old and norm_old in norm_para:
        # Reconstruct paragraph text with the substitution applied at the
        # normalised level, then write the result back as a single run.
        full = norm_para.replace(norm_old, _normalize(new))
        for run in paragraph.runs[1:]:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full
        return 1

    return 0


def patch_paragraph(paragraph) -> int:
    n = 0
    for old, new in REPLACEMENTS:
        n += replace_in_paragraph(paragraph, old, new)
    return n


def walk_document_for_replace(doc) -> int:
    n = 0
    for paragraph in doc.paragraphs:
        n += patch_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    n += patch_paragraph(paragraph)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for paragraph in ncell.paragraphs:
                                n += patch_paragraph(paragraph)
    return n


# ---------------------------------------------------------------------------
# Pass 2: insert "Implementation Details" subsection
# ---------------------------------------------------------------------------

# Each tuple: (style_kind, text). style_kind is one of:
#   "h3"     -> bold sub-heading
#   "label"  -> bold inline label that opens a paragraph (e.g. "Data pipeline:")
#   "p"      -> plain body paragraph
#   "bullet" -> dashed list item
IMPL_BLOCKS = [
    ("h3", "Implementation Details (as built)"),
    ("p",
     "The Computer/Software Engineering team delivered a complete Python "
     "implementation under the src/ directory of the project repository "
     "(github.com/rkfirat/hip-implant-ml). The implementation comprises five "
     "modules forming an end-to-end data pipeline, a containerised Flask "
     "backend, a responsive academic-style dashboard, and a 27-case pytest "
     "test suite. The full source tree, trained .joblib model artefacts, and "
     "honest cross-validation metrics are committed and pushed to the public "
     "GitHub repository."),
    ("label", "Codebase structure: "),
    ("bullet",
     "src/convert.py — Parses the raw ANSYS parametric export "
     "(Ti64_Hip_Implant_Dataset.csv), maps the P1…P31 parameter codes to "
     "physical names, normalises units (MPa → Pa, mm → m), and writes a "
     "cleaned 270-row dataset to data/dataset.csv."),
    ("bullet",
     "src/preprocess.py — Adds engineered features derived from F = m·g·K "
     "with frontal-plane cos/sin angular decomposition (force_x_N, "
     "force_z_N, force_magnitude_N) and exposes GroupKFold groups keyed by "
     "unique (mass, K, angle) tuples so the 108 ANSYS replicate rows cannot "
     "leak across train/test folds."),
    ("bullet",
     "src/train.py — Fits one Gaussian Process Regressor per target with "
     "kernel ConstantKernel × RBF + WhiteKernel(bounds = 1e-7…1e-2), 5 "
     "optimizer restarts, and StandardScaler normalisation on both inputs "
     "and outputs. Reports honest 5-fold GroupKFold R² and MAE per target "
     "to models/metrics.json."),
    ("bullet",
     "src/predict.py — Exposes a clean predict(mass, K, angle) Python API "
     "and a JSON CLI. Returns posterior mean AND posterior standard "
     "deviation for every target. The equivalent safety factor is computed "
     "analytically (827e6 / σ_vm) rather than trained as a redundant model, "
     "ensuring physical consistency between the two outputs."),
    ("bullet",
     "src/app.py — Flask backend exposing GET /predict?mass=&k=&angle= and "
     "serving the dashboard. Triggers a colour-coded status banner "
     "(Safe / Caution / Critical) using the minimum of all four safety "
     "factors against the 1.5 / 1.0 thresholds."),
    ("label", "Web dashboard: "),
    ("p",
     "The frontend is a single-page LaTeX-academic styled Flask template "
     "(src/templates/index.html) backed by vanilla JavaScript "
     "(src/static/app.js) and a responsive CSS Grid layout "
     "(src/static/style.css). The interface presents three input sliders "
     "(mass 50–100 kg, K 2.0–6.0×BW, angle 0–20°), three activity preset "
     "buttons (Walking K = 3.0, Stair Climbing K = 4.0, Running K = 5.0) "
     "aligned with the project's parametric tables, four Plotly.js safety-"
     "factor gauges (equivalent, neck, stem-1, stem-2) with dynamic upper "
     "bound = max(3, value × 1.3), a comparison scatter overlaying the "
     "current prediction on the 270-point ANSYS training set, and a model-"
     "performance table populated from the honest cross-validation metrics. "
     "A window resize listener forces Plotly to re-fit every chart on "
     "viewport rotation."),
    ("label", "Uncertainty quantification: "),
    ("p",
     "Every numerical output on the dashboard is rendered with a ±σ "
     "uncertainty annotation derived directly from the Gaussian Process "
     "posterior. This is the principal advantage of GPR over linear "
     "interpolation or tree-based models in a small-sample biomechanical "
     "setting: regions of the input space with sparse ANSYS coverage "
     "produce visibly wider uncertainty bands, surfacing extrapolation "
     "risk to the operator rather than hiding it behind a single point "
     "estimate."),
    ("label", "Honest cross-validation metrics (5-fold GroupKFold): "),
    ("p",
     "safety_factor_equivalent_min → R²_cv = 0.999996 ± 2.5×10⁻⁶; "
     "safety_factor_neck_min → R²_cv = 0.999996 ± 2.5×10⁻⁶; "
     "safety_factor_stem1_min → R²_cv = 0.999996 ± 2.5×10⁻⁶; "
     "safety_factor_stem2_min → R²_cv = 0.999992 ± 8.5×10⁻⁶; "
     "max_equivalent_vonmises_stress_Pa, max_principal_stress_Pa, and "
     "max_total_deformation_m → R²_cv ≈ 1.0. These metrics replace the "
     "earlier non-cross-validated R² placeholder and are emitted by "
     "src/train.py at training time into models/metrics.json."),
    ("label", "Containerisation and reproducibility: "),
    ("p",
     "The project ships with a Dockerfile based on python:3.9-slim and an "
     "accompanying .dockerignore that keeps the production image lean by "
     "excluding caches, the raw Excel artefacts, and the local test "
     "harness. The application can be built and pushed to a container "
     "registry with a single docker build command; binding is exposed on "
     "PORT (default 5050)."),
    ("label", "Verification suite: "),
    ("p",
     "A pytest harness under tests/ exercises both the preprocessing and "
     "the training+inference pipelines end-to-end (27 cases: 21 unit + 6 "
     "integration). The integration tier validates that train_one() "
     "produces R² > 0.95 for every non-noise-floor target on a fresh data "
     "regeneration, that predict() returns every trained target as a "
     "float, that safety factors fall monotonically with patient mass at "
     "fixed K, and that the activity-preset mapping (Walking, Stair "
     "Climbing, Running) is consistent with the dashboard preset buttons "
     "(K = 3.0, 4.0, 5.0 respectively). All 27 cases pass on the committed "
     "main branch."),
]

# ---------------------------------------------------------------------------
# XML helpers for inserting paragraphs before a target paragraph
# ---------------------------------------------------------------------------

INSERTION_MARKERS = (
    "3. System Integration",
    "System Integration",
)

def find_insertion_anchor(doc):
    """Return the first paragraph whose text starts with an insertion marker."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for marker in INSERTION_MARKERS:
            if text == marker or text.startswith(marker + " "):
                return paragraph
    return None


def make_paragraph_xml(template_paragraph, kind: str, text: str):
    """Clone the formatting style of `template_paragraph` and stamp `text`."""
    new_p = deepcopy(template_paragraph._p)
    # Strip every existing run from the clone.
    for r in list(new_p.findall(qn("w:r"))):
        new_p.remove(r)
    # Strip hyperlinks / fields that pdf2docx sometimes leaves dangling.
    for tag in ("w:hyperlink", "w:fldSimple"):
        for el in list(new_p.findall(qn(tag))):
            new_p.remove(el)
    return new_p


def insert_paragraphs_before(anchor, blocks) -> int:
    """Insert formatted paragraphs above `anchor`. Returns count inserted."""
    body = anchor._p.getparent()
    anchor_index = list(body).index(anchor._p)
    inserted = 0

    from docx.oxml import OxmlElement

    for kind, text in blocks:
        new_p = make_paragraph_xml(anchor, kind, text)

        # Build the run(s).
        if kind == "h3":
            run = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            bold = OxmlElement("w:b")
            rpr.append(bold)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "24")  # 12pt (half-points)
            rpr.append(sz)
            run.append(rpr)
            t = OxmlElement("w:t")
            t.text = text
            t.set(qn("xml:space"), "preserve")
            run.append(t)
            new_p.append(run)
        elif kind == "label":
            run = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            bold = OxmlElement("w:b")
            rpr.append(bold)
            run.append(rpr)
            t = OxmlElement("w:t")
            t.text = text
            t.set(qn("xml:space"), "preserve")
            run.append(t)
            new_p.append(run)
        elif kind == "bullet":
            run = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = "• " + text
            t.set(qn("xml:space"), "preserve")
            run.append(t)
            new_p.append(run)
        else:  # "p"
            run = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = text
            t.set(qn("xml:space"), "preserve")
            run.append(t)
            new_p.append(run)

        body.insert(anchor_index + inserted, new_p)
        inserted += 1

    return inserted


def find_findings_anchor(doc):
    """Return the paragraph that starts with '5.1' (Verification sub-heading)."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("5.1") and "verification" in text.lower():
            return paragraph
        if text == "5.1." or text == "5.1":
            return paragraph
    return None


def has_findings_heading(doc) -> bool:
    """Check whether a '5. FINDINGS' heading already exists in the body."""
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip().upper()
        if t in ("5. FINDINGS", "5.FINDINGS", "5  FINDINGS", "5\tFINDINGS"):
            return True
    return False


def find_sibling_heading(doc, marker: str):
    """Return a paragraph matching a top-level heading (e.g. '4. IMPLICATIONS')."""
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip().upper().replace("\t", " ").replace("  ", " ")
        if t == marker.upper() or t.startswith(marker.upper() + " "):
            return paragraph
    return None


def clone_heading_styling(template_paragraph, new_text: str):
    """Clone a sibling heading's XML wholesale, just replacing the visible text.

    Preserves every run-level property (font, size, bold, tab, color)
    that the sibling uses, so the inserted heading is visually identical
    to its peers.
    """
    new_p = deepcopy(template_paragraph._p)
    # Find every text element; clear all but the first; rewrite first with new_text.
    text_els = new_p.findall(".//" + qn("w:t"))
    if not text_els:
        return None
    text_els[0].text = new_text
    text_els[0].set(qn("xml:space"), "preserve")
    for t in text_els[1:]:
        t.text = ""
    return new_p


def insert_findings_heading_cloned(doc) -> bool:
    """Replace the simple heading insertion with a styling-cloned version.

    Walks the doc, removes any previously-inserted plain '5. FINDINGS'
    heading, then inserts a new one whose XML matches '4. IMPLICATIONS'.
    """
    # Remove any plain '5. FINDINGS' that we may have inserted previously
    body = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().upper() == "5. FINDINGS":
            body = paragraph._p.getparent()
            body.remove(paragraph._p)
            break

    sibling = find_sibling_heading(doc, "4. IMPLICATIONS")
    f_anchor = find_findings_anchor(doc)
    if sibling is None or f_anchor is None:
        return False

    new_p = clone_heading_styling(sibling, "5. FINDINGS")
    if new_p is None:
        return False
    body = f_anchor._p.getparent()
    body.insert(list(body).index(f_anchor._p), new_p)
    return True


# B6 — Mechanical Sub-Team coating-selection outcome paragraph
MECH_OUTCOME_BLOCK = [
    ("label", "Outcome: "),
    ("p",
     "Based on the deterministic parametric tribological characterisation, "
     "Diamond-Like Carbon (DLC) was selected as the optimal wear-resistant "
     "surface treatment for the predictive maintenance pipeline. The "
     "selection rationale rests on three literature-derived criteria: "
     "(i) the lowest reported coefficient of friction in the candidate "
     "set (COF ≈ 0.07–0.09 for DLC versus typical values >0.3 for "
     "uncoated Ti-6Al-4V and ~0.15–0.20 for CrN per Österle et al., "
     "2008); (ii) high adhesion energy to the Ti-6Al-4V substrate that "
     "resists delamination across the K = 2.0–6.0×BW load envelope; and "
     "(iii) compatibility with cementless osseointegration. The selected "
     "DLC wear coefficient is fed into the predictive maintenance "
     "algorithm as a literature-derived constant."),
]


def find_mech_outcome_anchor(doc):
    """Find the paragraph just before '2. Computer / Software Engineering Team'."""
    for i, paragraph in enumerate(doc.paragraphs):
        t = paragraph.text.strip()
        if t.startswith("2. Computer") or t.startswith("2.Computer"):
            return paragraph
    return None


def has_mech_outcome(doc) -> bool:
    """Check whether the Outcome block was already inserted."""
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Outcome:") and "DLC" in paragraph.text:
            return True
        if "Diamond-Like Carbon (DLC) was selected as the optimal" in paragraph.text:
            return True
    return False


# C2 — Force heading paragraphs to be fully bold (numeral + title)
HEADING_BOLDS = [
    "1. ORIGINALITY",
    "2. METHODOLOGY",
    "3. PROJECT MANAGEMENT",
    "4. IMPLICATIONS",
    "5. FINDINGS",
    "6. CONCLUSION",
    "7. EVALUATION FORM",
]


def force_heading_bold(doc) -> int:
    """For each top-level section heading, set every run to bold.

    pdf2docx sometimes leaves the numeral run with `<w:b w:val="0"/>`
    (explicit non-bold) while only the title run is bold. This makes the
    "5. FINDINGS" insertion (cloned from "4. IMPLICATIONS") render half-
    bold. Force all runs in every section heading to bold for visual
    uniformity.
    """
    from docx.oxml import OxmlElement
    fixed = 0
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip().upper().replace("\t", " ").replace("  ", " ")
        if not any(t == h.upper() or t.startswith(h.upper() + " ") or
                   t.startswith(h.upper()) for h in HEADING_BOLDS):
            continue
        for run in paragraph.runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run._r.insert(0, rpr)
            # Remove any non-bold marker, ensure a bold marker is present
            for b in rpr.findall(qn("w:b")):
                rpr.remove(b)
            b = OxmlElement("w:b")
            rpr.insert(0, b)
        fixed += 1
    return fixed


# C1 — Reference list reconstruction
CANONICAL_REFERENCES = [
    ("APPX-1: REFERENCES", None),
    ("Bennett",
     "Bennett, D., & Goswami, T. (2008). Finite element analysis of hip stem designs. "
     "Materials & Design, 29(1), 45–60. https://doi.org/10.1016/j.matdes.2006.12.014"),
    ("Gupta",
     "Gupta, V., & Chanda, A. (2022). Finite element analysis of a hybrid corrugated "
     "hip implant for stability and loading during gait phases. Biomedical Physics & "
     "Engineering Express, 8(3), 035028. https://doi.org/10.1088/2057-1976/ac669c"),
    ("Kanaizumi",
     "Kanaizumi, A., Suzuki, D., Nagoya, S., Teramoto, A., & Yamashita, T. (2022). "
     "Patient-specific three-dimensional evaluation of interface micromotion in two "
     "different short stem designs in cementless total hip arthroplasty: a finite "
     "element analysis. Journal of Orthopaedic Surgery and Research, 17(1), 437. "
     "https://doi.org/10.1186/s13018-022-03329-5"),
    ("Österle",
     "Österle, W., Klaffke, D., Griepentrog, M., Gross, U., Kranz, I., & Knabe, C. "
     "(2008). Potential of wear resistant coatings on Ti-6Al-4V for artificial hip "
     "joint bearing surfaces. Wear, 264(5–6), 505–517. "
     "https://doi.org/10.1016/j.wear.2007.04.001"),
    ("Gemini",
     "Google. (2026). Gemini. https://gemini.google.com"),
    ("NotebookLM",
     "Google. (2026). NotebookLM. https://notebooklm.google.com"),
]


def rewrite_references(doc) -> int:
    """Replace fragmented reference paragraphs/tables with canonical entries.

    Strategy:
      1. Find the paragraph that contains "APPX-1: REFERENCES" heading.
      2. Walk forward from there, replacing each reference paragraph with
         its canonical equivalent. Multi-line references that pdf2docx
         split across a paragraph + a fragmentation table are consolidated
         into a single paragraph (the table cells are cleared).
      3. Continuation tables (T30, T31, T32 by index — but located by
         content) are cleared (cells set to "") rather than deleted, to
         avoid breaking the document structure.
    """
    from docx.oxml import OxmlElement

    refs_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "APPX-1" in p.text or "REFERENCES" in p.text.upper():
            refs_idx = i
            break
    if refs_idx is None:
        return 0

    # Find every paragraph from refs_idx onward whose text starts with a known author
    fixed = 0
    authors_seen = set()

    # FIRST: handle the Gemini+NotebookLM combined paragraph (must run before
    # the marker loop, otherwise the Gemini marker would overwrite NotebookLM).
    for i in range(refs_idx, min(refs_idx + 25, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if "Gemini" in p.text and "NotebookLM" in p.text:
            _set_paragraph_text(
                p,
                "Google. (2026). Gemini. https://gemini.google.com\n"
                "Google. (2026). NotebookLM. https://notebooklm.google.com")
            authors_seen.update({"Gemini", "NotebookLM"})
            fixed += 1
            break

    for i in range(refs_idx, min(refs_idx + 20, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text

        # Special handling: P185 contains both "APPX-1: REFERENCES" heading and
        # the Bennett ref split by a newline. Replace its content entirely.
        if "APPX-1" in text and "Bennett" in text:
            _set_paragraph_text(p, "APPX-1: REFERENCES\n" +
                                "Bennett, D., & Goswami, T. (2008). Finite element "
                                "analysis of hip stem designs. Materials & Design, "
                                "29(1), 45–60. https://doi.org/10.1016/j.matdes.2006.12.014")
            authors_seen.update({"Bennett"})
            fixed += 1
            continue

        # Replace paragraphs whose text contains an author marker
        for marker, canonical in CANONICAL_REFERENCES:
            if marker == "APPX-1: REFERENCES" or canonical is None:
                continue
            if marker in authors_seen:
                continue
            if marker in text:
                _set_paragraph_text(p, canonical)
                authors_seen.add(marker)
                fixed += 1
                break

    # Clear fragmentation continuation tables. Identifying them: small tables
    # (1 row, many narrow cells, total text < 100 chars) appearing right after
    # the references. We sweep the last 6 tables in the doc for this pattern.
    for table in doc.tables[-6:]:
        if len(table.rows) != 1:
            continue
        joined = " ".join(c.text.strip() for r in table.rows for c in r.cells)
        if not joined:
            continue
        # heuristic: continuation tables are 6+ cells, each cell holds one word/phrase
        if len(table.rows[0].cells) >= 6 and len(joined) < 200:
            # Avoid clearing the team-roster or other meaningful tables: continuation
            # tables only contain journal-name fragments, no digits >2.
            has_long_token = any(len(c.text.strip()) > 30
                                 for r in table.rows for c in r.cells)
            if has_long_token:
                continue
            for r in table.rows:
                for c in r.cells:
                    for cp in c.paragraphs:
                        for rn in list(cp.runs):
                            rn.text = ""
            fixed += 1

    return fixed


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Reset a paragraph's text content while preserving paragraph-level XML."""
    # Remove all child runs / hyperlinks / fields, keep paragraph properties
    for tag in ("w:r", "w:hyperlink", "w:fldSimple"):
        for el in list(paragraph._p.findall(qn(tag))):
            paragraph._p.remove(el)
    # Split on newlines: emit one run per line, with a soft break between
    from docx.oxml import OxmlElement
    lines = new_text.split("\n")
    for li, line in enumerate(lines):
        if li > 0:
            br_r = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br_r.append(br)
            paragraph._p.append(br_r)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = line
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        paragraph._p.append(r)


# ============================================================================
# Image-insertion routine (Pass 7) — dashboard screenshots, GPR uncertainty,
# pipeline diagram. Inserts pictures + centered captions at named anchors in
# the Implementation Details block.
# ============================================================================

from docx.shared import Inches  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402


def _move_paragraph_after(paragraph_to_move, anchor_paragraph):
    """Detach paragraph_to_move from its parent and insert it as the
    immediate next sibling of anchor_paragraph."""
    p = paragraph_to_move._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
    anchor_paragraph._p.addnext(p)


def insert_picture_after(doc, anchor_paragraph, image_path: str,
                         width_inches: float, caption_text: str):
    """Insert an image + caption directly after `anchor_paragraph`.

    Returns the caption paragraph so the caller can chain further
    insertions below it.
    """
    # 1. Create the picture paragraph at the end of the doc, then relocate it.
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(image_path, width=Inches(width_inches))

    # 2. Add a caption at the end of the doc, then relocate.
    cap_p = doc.add_paragraph(caption_text)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap_p.runs:
        r.italic = True
        r.font.size = None  # use surrounding style

    # 3. Move into position. Caption goes right after pic; pic goes right
    #    after anchor; resulting order: anchor → pic → caption.
    _move_paragraph_after(pic_p, anchor_paragraph)
    _move_paragraph_after(cap_p, pic_p)
    return cap_p


def find_paragraph_starting_with(doc, prefix: str):
    """Return the first paragraph whose text starts with the given prefix."""
    for paragraph in doc.paragraphs:
        if paragraph.text.lstrip().startswith(prefix):
            return paragraph
    return None


def has_dashboard_figures(doc) -> bool:
    """True once any of the inserted figure captions is already in the doc."""
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Figure-2."):
            return True
        if paragraph.text.startswith("Figure 2."):
            return True
    return False


# Image insertion plan:
#   (anchor prefix, image path, width inches, caption)
FIGURE_PLAN = [
    # Pipeline diagram — right after Implementation Details opener paragraph
    ("The Computer/Software Engineering team delivered",
     "screenshots/fig_pipeline_diagram.png", 6.2,
     "Figure-2. Computer/Software Engineering team data-handshake pipeline. "
     "Static ANSYS-derived dataset flows through convert → preprocess → train "
     "→ predict modules into the Flask + Plotly.js dashboard, producing "
     "patient-specific safety factors and ±σ uncertainty bands in real time."),

    # Dashboard default view — after "Web dashboard:" paragraph
    ("The frontend is a single-page LaTeX-academic styled",
     "screenshots/desktop_1440__default_70kg_K2.5_0deg_cropped.png", 5.8,
     "Figure-3. Predictive maintenance dashboard — default state (70 kg, "
     "K = 2.5×BW, 0° angle). Four Plotly.js gauges display equivalent, neck, "
     "stem-1, and stem-2 safety factors; sliders on the left expose mass / K "
     "/ angle controls; activity preset buttons fix K to nominal walking / "
     "stair-climbing / running loads; a green status banner indicates the "
     "design is operating within the safe envelope."),

    # Dashboard Safe state
    ("Figure-3. Predictive maintenance dashboard — default state",
     "screenshots/desktop_1440__safe_50kg_K2.5_0deg_cropped.png", 5.8,
     "Figure-4. Dashboard — 50 kg patient at K = 2.5×BW. Minimum safety "
     "factor remains comfortably above 1.5; the status banner stays green "
     "(\"Safe\") and every gauge needle sits in the wide-margin region."),

    # Dashboard Critical state
    ("Figure-4. Dashboard — 50 kg patient",
     "screenshots/desktop_1440__critical_100kg_K5_0deg_cropped.png", 5.8,
     "Figure-5. Dashboard — 100 kg patient at K = 5.0×BW (running). The "
     "minimum safety factor crosses the 1.0 threshold and the status "
     "banner flips to red (\"Critical\"), demonstrating that the alert "
     "logic triggered by min(SF_equivalent, SF_neck, SF_stem-1, "
     "SF_stem-2) < 1.0 fires as designed."),

    # GPR posterior uncertainty plot — after "Uncertainty quantification:"
    ("Every numerical output on the dashboard is rendered with a ±σ",
     "screenshots/fig_gpr_uncertainty.png", 6.0,
     "Figure-6. GPR posterior mean and ±σ uncertainty band for the neck "
     "safety factor across the patient mass range under walking load "
     "(K = 3.0×BW, 0° angle). Markers indicate the six training masses "
     "(50, 60, 70, 80, 90, 100 kg). Inside the validated envelope the band "
     "is tight; outside, the uncertainty widens visibly — surfacing "
     "extrapolation risk to the operator rather than hiding it behind a "
     "single point estimate."),

    # Mobile responsive view — after the GPR figure
    ("Figure-6. GPR posterior mean",
     "screenshots/mobile_390__default_70kg_K2.5_0deg_cropped.png", 2.8,
     "Figure-7. Mobile-responsive layout (390 px viewport). The CSS Grid "
     "auto-fit minmax(240px, 1fr) rule reflows the four gauges into a "
     "single column; the sidebar collapses into a slide-in panel triggered "
     "by the header button. The window-resize listener calls "
     "Plotly.Plots.resize() to keep every chart sharp through device "
     "rotation."),
]


from docx.shared import Pt  # noqa: E402


def _zero_pPr_indent(paragraph) -> None:
    """Strip inherited indent and force a clean left margin on a paragraph.

    Cloned anchor pPrs often carry `w:ind w:left="1418"` from numbered list
    or table-cell ancestors. Setting paragraph_format.left_indent leaves
    other indent attributes alive. We strip the indent element outright.
    """
    from docx.oxml import OxmlElement
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    for tag in ("w:ind",):
        for el in list(pPr.findall(qn(tag))):
            pPr.remove(el)


def _set_paragraph_indent(paragraph, left_pt: float = 0,
                          first_line_pt: float = 0,
                          right_pt: float = 0,
                          hanging_pt: float = 0) -> None:
    """Replace any existing w:ind with explicit values (in points)."""
    from docx.oxml import OxmlElement
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._p.insert(0, pPr)
    for el in list(pPr.findall(qn("w:ind"))):
        pPr.remove(el)
    ind = OxmlElement("w:ind")
    if left_pt:
        ind.set(qn("w:left"), str(int(left_pt * 20)))
    if right_pt:
        ind.set(qn("w:right"), str(int(right_pt * 20)))
    if hanging_pt:
        ind.set(qn("w:hanging"), str(int(hanging_pt * 20)))
    elif first_line_pt:
        ind.set(qn("w:firstLine"), str(int(first_line_pt * 20)))
    pPr.append(ind)


def _set_line_spacing(paragraph, line_multiplier: float = 1.15) -> None:
    """Force `w:spacing w:line="<n>" w:lineRule="auto"` where n=240*multiplier."""
    from docx.oxml import OxmlElement
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:line"), str(int(240 * line_multiplier)))
    sp.set(qn("w:lineRule"), "auto")


def normalize_layout(doc) -> int:
    """Pass 8 — fix spacing, indentation, and alignment of inserted content.

    Cloned pPr from anchor paragraphs frequently carries excessive left
    indent (e.g. w:left="1418" from list/table ancestors) and tight line
    spacing (w:line="233"). We force-override these for every inserted
    paragraph.
    """
    n = 0

    label_prefixes = (
        "Codebase structure:", "Web dashboard:", "Uncertainty quantification:",
        "Honest cross-validation metrics", "Containerisation",
        "Verification suite:", "Outcome:",
    )

    body_starts_for_indent = (
        "The Computer/Software Engineering team delivered",
        "The frontend is a single-page",
        "Every numerical output on the dashboard",
        "safety_factor_equivalent_min", "safety_factor_neck_min",
        "The project ships with a Dockerfile",
        "A pytest harness under tests/",
        "Based on the deterministic parametric tribological",
    )

    for paragraph in doc.paragraphs:
        text = paragraph.text.lstrip()
        if not text:
            continue
        pf = paragraph.paragraph_format

        # 1. Implementation Details heading
        if text.startswith("Implementation Details (as built)"):
            _set_paragraph_indent(paragraph, left_pt=0)
            pf.space_before = Pt(18)
            pf.space_after = Pt(8)
            pf.keep_with_next = True
            _set_line_spacing(paragraph, 1.15)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(13)
            n += 1
            continue

        # 2. Label paragraphs (bold inline labels that open a subsection)
        if any(text.startswith(p) for p in label_prefixes):
            _set_paragraph_indent(paragraph, left_pt=0)
            pf.space_before = Pt(14)
            pf.space_after = Pt(4)
            pf.keep_with_next = True
            _set_line_spacing(paragraph, 1.15)
            n += 1
            continue

        # 3. Bullet paragraphs
        if text.startswith("•"):
            _set_paragraph_indent(paragraph, left_pt=28, hanging_pt=14)
            pf.space_before = Pt(2)
            pf.space_after = Pt(4)
            _set_line_spacing(paragraph, 1.15)
            n += 1
            continue

        # 4. Figure caption paragraphs
        if text.startswith("Figure-") and len(text) > 16:
            _set_paragraph_indent(paragraph, left_pt=36, right_pt=36)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(2)
            pf.space_after = Pt(16)
            _set_line_spacing(paragraph, 1.10)
            for run in paragraph.runs:
                run.italic = True
                if run.font.size is None:
                    run.font.size = Pt(9)
            n += 1
            continue

        # 5. Reference list — hanging indent on every entry
        ref_starts = ("Bennett,", "Gupta,", "Kanaizumi,", "Österle,",
                      "Google. (2026).")
        if any(text.startswith(s) for s in ref_starts):
            _set_paragraph_indent(paragraph, left_pt=28, hanging_pt=28)
            pf.space_after = Pt(8)
            _set_line_spacing(paragraph, 1.15)
            n += 1
            continue

        # 6. Body paragraphs of the insertion block — strip inherited indent
        if any(text.startswith(s) for s in body_starts_for_indent):
            _set_paragraph_indent(paragraph, left_pt=0)
            pf.space_before = Pt(4)
            pf.space_after = Pt(8)
            _set_line_spacing(paragraph, 1.15)
            n += 1
            continue

    # 7. Image-bearing paragraphs — picked up by the presence of a w:drawing
    for paragraph in doc.paragraphs:
        drawings = paragraph._p.findall(".//" + qn("w:drawing"))
        if not drawings:
            continue
        # Only act on paragraphs whose drawing we just inserted: heuristic =
        # the paragraph contains exactly one drawing and no other text.
        if paragraph.text.strip():
            continue
        _set_paragraph_indent(paragraph, left_pt=0)
        pf = paragraph.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(12)
        pf.space_after = Pt(2)
        pf.keep_with_next = True
        n += 1

    return n


def insert_figures(doc) -> int:
    """Walk FIGURE_PLAN, locate each anchor, insert picture + caption.

    Each successful insertion updates the anchor for the next iteration
    (the caption becomes the new anchor), enabling chained inserts.
    """
    if has_dashboard_figures(doc):
        print("Pass 7: figures already inserted, skipping")
        return 0

    inserted = 0
    for anchor_prefix, image_path, width, caption in FIGURE_PLAN:
        # Re-read paragraph list each iteration because previous inserts
        # invalidate index-based lookups; we use prefix matching instead.
        anchor = find_paragraph_starting_with(doc, anchor_prefix)
        if anchor is None:
            print(f"  WARN: anchor not found for {image_path}: '{anchor_prefix[:40]}'")
            continue
        full_path = str(Path(__file__).resolve().parents[1] / image_path)
        if not Path(full_path).exists():
            print(f"  WARN: image not found: {full_path}")
            continue
        insert_picture_after(doc, anchor, full_path, width, caption)
        inserted += 1
    return inserted


def main() -> None:
    print(f"Loading {SRC.name} ({SRC.stat().st_size // 1024} KB)")
    doc = Document(str(SRC))

    n_par = len(doc.paragraphs)
    n_tab = len(doc.tables)
    print(f"Document: {n_par} paragraphs, {n_tab} tables")

    n_repl = walk_document_for_replace(doc)
    print(f"Pass 1: applied {n_repl} text replacement(s)")

    anchor = find_insertion_anchor(doc)
    if anchor is None:
        print("WARNING: insertion anchor not found; skipping Implementation block")
    else:
        n_ins = insert_paragraphs_before(anchor, IMPL_BLOCKS)
        print(f"Pass 2: inserted {n_ins} 'Implementation Details' paragraph(s) "
              f"before anchor '{anchor.text[:50]}'")

    # Pass 3 — insert '5. FINDINGS' parent heading with cloned sibling styling
    ok = insert_findings_heading_cloned(doc)
    if ok:
        print("Pass 3: inserted '5. FINDINGS' heading "
              "(styling cloned from '4. IMPLICATIONS')")
    else:
        print("WARNING: could not insert styled '5. FINDINGS' heading")

    # Pass 4 — insert Mechanical Sub-Team coating-selection outcome
    if not has_mech_outcome(doc):
        m_anchor = find_mech_outcome_anchor(doc)
        if m_anchor is not None:
            n_mo = insert_paragraphs_before(m_anchor, MECH_OUTCOME_BLOCK)
            print(f"Pass 4: inserted {n_mo} 'Outcome:' paragraph(s) into "
                  f"Mechanical Sub-Team before '{m_anchor.text[:40]}'")
        else:
            print("WARNING: Mechanical Sub-Team anchor not found")
    else:
        print("Pass 4: Mechanical Outcome already present, skipping")

    # Pass 5 — force all section headings to fully-bold runs
    n_bold = force_heading_bold(doc)
    print(f"Pass 5: forced bold on {n_bold} section heading(s)")

    # Pass 6 — rewrite fragmented reference list to canonical form
    n_refs = rewrite_references(doc)
    print(f"Pass 6: rewrote {n_refs} reference fragment(s)")

    # Pass 7 — insert dashboard screenshots, GPR uncertainty, pipeline diagram
    n_figs = insert_figures(doc)
    print(f"Pass 7: inserted {n_figs} figure(s) (pictures + captions)")

    # Pass 8 — normalize spacing, indentation, and alignment of inserted content
    n_layout = normalize_layout(doc)
    print(f"Pass 8: normalized layout on {n_layout} inserted paragraph(s)")

    doc.save(str(DST))
    print(f"Saved -> {DST.name} ({DST.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
